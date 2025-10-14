# app.py
# Streamlit: Excel/CSV -> (BEX / Non-BEX) Review/Plan .docx (ZIP)

import io
import re
import zipfile
from typing import Any, Dict, Optional

import pandas as pd
import streamlit as st
from docx import Document
from docx.oxml.ns import qn

# ───────────────────────────── UI CONFIG ─────────────────────────────
st.set_page_config(page_title="Excel → Review/Plan Generator", layout="wide")
st.title("📊 Excel/CSV → 📄 Review/Plan Generator (BEX & Non-BEX)")

# ───────────────────────────── HELPERS ─────────────────────────────
PLACEHOLDERS = [
    "title", "plan_month", "store", "bex",
    "plan_vs_target", "mobile_actual", "mobile_target",
    "fixed_target", "fixed_actual",
    "voice_vs_target", "fixed_vs_target",
    "llu_actual", "nga_actual", "ftth_actual",
    "eon_tv_actual", "fwa_actual",
    "mobile_upgrades", "fixed_upgrades",
    "pending_mobile", "pending_fixed",
]

BEX_SET_DEFAULT = {"DRZ01", "FKM01", "ESC01", "LND01", "PKK01"}


def normkey(x: str) -> str:
    return re.sub(r"[\s\-_\.]+", "", str(x).strip().lower())


def a1_to_idx(letter: str) -> int:
    """Convert Excel column letters (A, Z, AA…) to 0-based index."""
    s = letter.strip().upper()
    if not s.isalpha():
        raise ValueError(f"Not a letter: {letter}")
    n = 0
    for ch in s:
        n = n * 26 + (ord(ch) - 64)
    return n - 1


def get_by_letter(row: pd.Series, letter: Optional[str]) -> Any:
    if not letter:
        return ""
    try:
        idx = a1_to_idx(letter)
        # row is a Series with same ordering as df columns:
        if idx < 0 or idx >= len(row.index):
            return ""
        v = row.iloc[idx]
        return "" if pd.isna(v) else v
    except Exception:
        return ""


def set_default_font(doc: Document, font_name: str = "Aptos") -> None:
    for style in doc.styles:
        if hasattr(style, "font"):
            try:
                style.font.name = font_name
                style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                style._element.rPr.rFonts.set(qn("w:cs"), font_name)
            except Exception:
                pass


def replace_placeholders(doc: Document, mapping: Dict[str, Any]) -> None:
    pat = re.compile(r"\[\[([A-Za-z0-9_]+)\]\]")

    def subfun(s: str) -> str:
        return pat.sub(lambda m: "" if mapping.get(m.group(1)) is None else str(mapping.get(m.group(1), "")), s)

    for p in doc.paragraphs:
        for r in p.runs:
            r.text = subfun(r.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = subfun(r.text)


def read_data(uploaded, sheet_name: str) -> Optional[pd.DataFrame]:
    """Reads CSV or XLSX. Returns DataFrame or None."""
    try:
        fname = getattr(uploaded, "name", "")
        if fname.lower().endswith(".csv"):
            df = pd.read_csv(uploaded)
        else:
            xfile = pd.ExcelFile(uploaded, engine="openpyxl")
            if sheet_name and sheet_name not in xfile.sheet_names:
                st.error(f"Το sheet '{sheet_name}' δεν βρέθηκε. Διαθέσιμα: {xfile.sheet_names}")
                return None
            sheet = sheet_name if sheet_name else xfile.sheet_names[0]
            df = pd.read_excel(xfile, sheet_name=sheet, engine="openpyxl")

        # ── PATCH: auto-fix duplicate headers ─────────────────────────
        orig_cols = list(df.columns)
        if pd.Series(orig_cols).duplicated(keep=False).any():
            def uniquify(seq):
                seen = {}
                out = []
                for x in seq:
                    n = seen.get(x, 0)
                    out.append(x if n == 0 else f"{x}__{n+1}")
                    seen[x] = n + 1
                return out

            new_cols = uniquify(orig_cols)
            df.columns = new_cols
            dups = sorted({c for c in orig_cols if orig_cols.count(c) > 1})
            st.warning(
                "Excel: Βρέθηκαν διπλότυπες κεφαλίδες και έγιναν auto-rename.\n\n"
                f"Διπλότυπα: {dups}\n\n"
                f"Νέα ονόματα (ενδεικτικά): {new_cols[:12]} ..."
            )
        # ──────────────────────────────────────────────────────────────

        return df
    except Exception as e:
        st.error(f"Δεν άνοιξε το αρχείο: {e}")
        return None


def fmt_percent(x: Any) -> Any:
    """If value looks like ratio (e.g., 1.22), turn to 122% string."""
    try:
        if x == "" or x is None:
            return ""
        xv = float(x)
        return f"{round(xv * 100)}%"
    except Exception:
        return x


# ───────────────────────────── SIDEBAR ─────────────────────────────
with st.sidebar:
    st.header("Ρυθμίσεις")
    debug_mode = st.toggle("🛠 Debug mode", value=True)
    test_mode = st.toggle("🧪 Test mode (πρώτες 50 γραμμές)", value=True)

    st.subheader("Templates (.docx)")
    tpl_bex = st.file_uploader("BEX template", type=["docx"], key="tpl1")
    tpl_nonbex = st.file_uploader("Non-BEX template", type=["docx"], key="tpl2")
    st.caption(
        "Placeholders στο Word: " +
        ", ".join([f"[[{k}]]" for k in PLACEHOLDERS])
    )

st.markdown("### 1) Ανέβασε Excel")
uploaded = st.file_uploader("Drag and drop file here", type=["xlsx", "csv"])
sheet_name = st.text_input("Όνομα φύλλου (Sheet - μόνο για Excel)", value="Sheet1")

with st.expander("📌 Ρύθμιση γραμμών (headers & δεδομένα)", expanded=False):
    st.write("Αν η 1η γραμμή είναι κεφαλίδες, άφησέ το ως έχει. Αλλιώς μπορείς να προσαρμόσεις.")
    start_row = st.number_input("Πρώτη γραμμή δεδομένων (1-based, default=2)", min_value=1, value=2, step=1)

with st.expander("🏬 STORE & BEX", expanded=True):
    store_options = st.radio("Πώς διαβάζουμε το STORE;", ["Από κεφαλίδα 'Dealer_Code' / 'Shop Code'", "Από γράμμα στήλης"], index=0)
    store_letter = ""
    if store_options == "Από γράμμα στήλης":
        store_letter = st.text_input("Γράμμα στήλης για STORE (π.χ. A)", value="A")

    bex_mode = st.radio("Πώς βρίσκουμε αν είναι BEX;", ["Λίστα κωδικών", "Από στήλη (YES/NO)"], index=0)
    bex_from_list = set()
    bex_yesno_letter = ""
    if bex_mode == "Λίστα κωδικών":
        bex_txt = st.text_area("BEXStores (comma-separated)", "DRZ01,FKM01,ESC01,LND01,PKK01")
        bex_from_list = {s.strip().upper() for s in bex_txt.split(",") if s.strip()}
    else:
        bex_yesno_letter = st.text_input("Γράμμα στήλης BEX (YES/NO)", value="J")

with st.expander("🔤 Mapping με γράμματα Excel (A, N, AA, AF, AH)", expanded=True):
    letter_plan_vs = st.text_input("plan_vs_target", value="A")
    letter_mobile_act = st.text_input("mobile_actual", value="N")
    letter_mobile_tgt = st.text_input("mobile_target", value="O")
    letter_fixed_tgt = st.text_input("fixed_target", value="P")
    letter_fixed_act = st.text_input("fixed_actual", value="Q")
    letter_voice_vs = st.text_input("voice_vs_target", value="R")
    letter_fixed_vs = st.text_input("fixed_vs_target", value="S")
    letter_llu = st.text_input("llu_actual", value="T")
    letter_nga = st.text_input("nga_actual", value="U")
    letter_ftth = st.text_input("ftth_actual", value="V")
    letter_eon = st.text_input("eon_tv_actual", value="X")
    letter_fwa = st.text_input("fwa_actual", value="Y")
    letter_mup = st.text_input("mobile_upgrades", value="AA")
    letter_fup = st.text_input("fixed_upgrades", value="AB")
    letter_pmob = st.text_input("pending_mobile", value="AF")
    letter_pfix = st.text_input("pending_fixed", value="AH")

st.markdown("### 2) Προεπισκόπηση τιμών που θα περάσουν")
preview_out = st.empty()

run = st.button("⚙️ Generate")

# ───────────────────────────── MAIN ─────────────────────────────
if run:
    # checks
    if not uploaded:
        st.error("Ανέβασε Excel/CSV πρώτα.")
        st.stop()
    if not tpl_bex or not tpl_nonbex:
        st.error("Ανέβασε και τα 2 templates (.docx).")
        st.stop()

    file_kb = len(uploaded.getbuffer()) / 1024
    st.info(f"📄 Δεδομένα: {file_kb:.1f} KB | BEX tpl: {tpl_bex.size/1024:.1f} KB | Non-BEX tpl: {tpl_nonbex.size/1024:.1f} KB")

    df = read_data(uploaded, sheet_name)
    if df is None or df.empty:
        st.error("Αποτυχία ανάγνωσης αρχείου ή άδειος πίνακας.")
        st.stop()

    # Αν ο χρήστης είπε ότι τα δεδομένα ξεκινούν από γραμμή >1, κόψε τις πρώτες (start_row-2) γραμμές,
    # ώστε το row index 0 να αντιστοιχεί στη γραμμή 'start_row'.
    if start_row > 1:
        df = df.iloc[start_row-2:].reset_index(drop=True)

    if debug_mode:
        st.caption(f"Columns ({len(df.columns)}): {list(df.columns)}")
        st.dataframe(df.head(10))

    # ── PREVIEW MAPPING (πρώτες 3 σειρές) ────────────────────────────
    preview_rows = []
    max_prev = min(3, len(df))
    for ridx in range(max_prev):
        row = df.iloc[ridx]
        # store
        if store_options == "Από γράμμα στήλης":
            store_val = str(get_by_letter(row, store_letter)).strip()
            store_from = f"letter {store_letter}"
        else:
            # Προσπαθώ από κλασικές κεφαλίδες:
            candidates = ["Dealer_Code", "Dealer Code", "Shop Code", "Shop_Code", "ShopCode"]
            store_val = ""
            for c in candidates:
                if c in df.columns:
                    v = row[c]
                    store_val = "" if pd.isna(v) else str(v)
                    if store_val:
                        break
            store_from = "header"
        # bex
        if bex_mode == "Λίστα κωδικών":
            bex_val = "YES" if store_val.upper() in bex_from_list else "NO"
        else:
            b = str(get_by_letter(row, bex_yesno_letter)).strip().lower()
            bex_val = "YES" if b in ("yes", "y", "1", "true", "ναι") else "NO"

        mapped = {
            "row_excel": start_row + ridx,
            "store": {"from": store_from, "value": store_val},
            "bex": bex_val,
            "plan_vs_target": get_by_letter(row, letter_plan_vs),
            "mobile_actual": get_by_letter(row, letter_mobile_act),
            "mobile_target": get_by_letter(row, letter_mobile_tgt),
            "fixed_target": get_by_letter(row, letter_fixed_tgt),
            "fixed_actual": get_by_letter(row, letter_fixed_act),
            "voice_vs_target": get_by_letter(row, letter_voice_vs),
            "fixed_vs_target": get_by_letter(row, letter_fixed_vs),
            "llu_actual": get_by_letter(row, letter_llu),
            "nga_actual": get_by_letter(row, letter_nga),
            "ftth_actual": get_by_letter(row, letter_ftth),
            "eon_tv_actual": get_by_letter(row, letter_eon),
            "fwa_actual": get_by_letter(row, letter_fwa),
            "mobile_upgrades": get_by_letter(row, letter_mup),
            "fixed_upgrades": get_by_letter(row, letter_fup),
            "pending_mobile": get_by_letter(row, letter_pmob),
            "pending_fixed": get_by_letter(row, letter_pfix),
        }
        preview_rows.append(mapped)

    preview_out.json(preview_rows)

    # ── BUILD DOCS ───────────────────────────────────────────────────
    tpl_bex_bytes = tpl_bex.read()
    tpl_nonbex_bytes = tpl_nonbex.read()

    out_zip = io.BytesIO()
    zf = zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED)
    built = 0
    total_rows = len(df) if not test_mode else min(50, len(df))
    pbar = st.progress(0, text="Δημιουργία εγγράφων…")

    for idx in range(total_rows):
        row = df.iloc[idx]
        # STORE
        if store_options == "Από γράμμα στήλης":
            store = str(get_by_letter(row, store_letter)).strip()
        else:
            store = ""
            for c in ["Dealer_Code", "Dealer Code", "Shop Code", "Shop_Code", "ShopCode"]:
                if c in df.columns:
                    v = row[c]
                    store = "" if pd.isna(v) else str(v)
                    if store:
                        break
        if not store:
            pbar.progress((idx + 1) / total_rows, text=f"Παράλειψη γραμμής {start_row + idx} (κενό store)")
            continue
        store_up = store.upper()

        # BEX
        if bex_mode == "Λίστα κωδικών":
            is_bex = store_up in bex_from_list
        else:
            b = str(get_by_letter(row, bex_yesno_letter)).strip().lower()
            is_bex = b in ("yes", "y", "1", "true", "ναι")

        # values
        mapping = {
            "title": f"Review September 2025 — Plan October 2025 — {store_up}",
            "plan_month": "Review September 2025 — Plan October 2025",
            "store": store_up,
            "bex": "YES" if is_bex else "NO",
            "plan_vs_target": get_by_letter(row, letter_plan_vs),
            "mobile_actual": get_by_letter(row, letter_mobile_act),
            "mobile_target": get_by_letter(row, letter_mobile_tgt),
            "fixed_target": get_by_letter(row, letter_fixed_tgt),
            "fixed_actual": get_by_letter(row, letter_fixed_act),
            "voice_vs_target": fmt_percent(get_by_letter(row, letter_voice_vs)),
            "fixed_vs_target": fmt_percent(get_by_letter(row, letter_fixed_vs)),
            "llu_actual": get_by_letter(row, letter_llu),
            "nga_actual": get_by_letter(row, letter_nga),
            "ftth_actual": get_by_letter(row, letter_ftth),
            "eon_tv_actual": get_by_letter(row, letter_eon),
            "fwa_actual": get_by_letter(row, letter_fwa),
            "mobile_upgrades": get_by_letter(row, letter_mup),
            "fixed_upgrades": get_by_letter(row, letter_fup),
            "pending_mobile": get_by_letter(row, letter_pmob),
            "pending_fixed": get_by_letter(row, letter_pfix),
        }

        # build doc
        try:
            doc = Document(io.BytesIO(tpl_bex_bytes if is_bex else tpl_nonbex_bytes))
            set_default_font(doc, "Aptos")
            replace_placeholders(doc, mapping)
            out_name = f"{store_up}_ReviewSep_PlanOct.docx"
            buf = io.BytesIO()
            doc.save(buf)
            zf.writestr(out_name, buf.getvalue())
            built += 1
            pbar.progress((idx + 1) / total_rows, text=f"Φτιάχνω: {out_name} ({idx + 1}/{total_rows})")
        except Exception as e:
            st.warning(f"⚠️ Σφάλμα στη γραμμή {start_row + idx}: {e}")
            if debug_mode:
                st.exception(e)

    zf.close()
    pbar.empty()

    if built == 0:
        st.error("Δεν δημιουργήθηκε αρχείο. Έλεγξε STORE mapping & templates.")
    else:
        st.success(f"Έτοιμα {built} αρχεία.")
        st.download_button("⬇️ Κατέβασε ZIP", data=out_zip.getvalue(), file_name="reviews_from_excel.zip")